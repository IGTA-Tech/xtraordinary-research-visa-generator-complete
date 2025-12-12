"""
File Processor - Extract text from uploaded files using Mistral Vision
"""
import io
import base64
from typing import Tuple
import logging

from app.config import settings

logger = logging.getLogger(__name__)


async def extract_with_mistral_vision(image_data: bytes, filename: str) -> str:
    """
    Use Mistral's Pixtral vision model to extract text from images/documents.

    Args:
        image_data: Image bytes (PNG, JPG, or PDF page rendered as image)
        filename: Original filename for context

    Returns:
        Extracted text content
    """
    try:
        from mistralai import Mistral

        client = Mistral(api_key=settings.MISTRAL_API_KEY)

        # Encode image to base64
        base64_image = base64.b64encode(image_data).decode("utf-8")

        # Determine mime type
        filename_lower = filename.lower()
        if filename_lower.endswith(".png"):
            mime_type = "image/png"
        elif filename_lower.endswith((".jpg", ".jpeg")):
            mime_type = "image/jpeg"
        elif filename_lower.endswith(".gif"):
            mime_type = "image/gif"
        elif filename_lower.endswith(".webp"):
            mime_type = "image/webp"
        else:
            mime_type = "image/png"  # Default for PDF pages

        # Call Mistral Pixtral for OCR/text extraction
        response = client.chat.complete(
            model="pixtral-12b-2409",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": f"data:{mime_type};base64,{base64_image}"
                        },
                        {
                            "type": "text",
                            "text": """Extract ALL text content from this document/image.

Instructions:
- Extract every piece of text you can see, preserving the structure
- Include headers, paragraphs, bullet points, table content
- Maintain the reading order (top to bottom, left to right)
- If this is a form or structured document, preserve the field labels and values
- Do not summarize - extract the complete text verbatim
- If there are multiple columns, process them in order

Return ONLY the extracted text, no commentary."""
                        }
                    ]
                }
            ],
            max_tokens=8192,
        )

        extracted_text = response.choices[0].message.content
        logger.info(f"Mistral extracted {len(extracted_text)} chars from {filename}")
        return extracted_text

    except Exception as e:
        logger.error(f"Mistral vision extraction failed: {e}")
        return ""


async def extract_text_from_pdf(content: bytes, filename: str) -> Tuple[str, int]:
    """
    Extract text from PDF - uses PyMuPDF for text-based PDFs,
    falls back to Mistral Vision for scanned/image PDFs.

    Returns:
        Tuple of (extracted_text, page_count)
    """
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        page_count = len(pdf)

        for page_num, page in enumerate(pdf):
            # First try regular text extraction
            page_text = page.get_text().strip()

            # If page has very little text, it might be scanned - use Mistral
            if len(page_text) < 100 and settings.MISTRAL_API_KEY:
                logger.info(f"Page {page_num + 1} appears scanned, using Mistral Vision")

                # Render page to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("png")

                # Extract with Mistral
                page_text = await extract_with_mistral_vision(img_bytes, f"{filename}_page_{page_num}.png")

            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

        pdf.close()

        text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(text)} chars from {page_count} page PDF")
        return text, page_count

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return "", 0


async def extract_text_from_docx(content: bytes) -> Tuple[str, int]:
    """
    Extract text from DOCX using python-docx.

    Returns:
        Tuple of (extracted_text, estimated_page_count)
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n\n".join(text_parts)

        # Estimate pages (roughly 500 words per page)
        word_count = len(text.split())
        page_count = max(1, word_count // 500)

        logger.info(f"Extracted {len(text)} chars from DOCX (~{page_count} pages)")
        return text, page_count

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return "", 0


async def extract_text_from_image(content: bytes, filename: str) -> Tuple[str, int]:
    """
    Extract text from image using Mistral Vision.

    Returns:
        Tuple of (extracted_text, 1)
    """
    if not settings.MISTRAL_API_KEY:
        logger.warning("MISTRAL_API_KEY not set - cannot process images")
        return "[Image content - Mistral API key required for OCR]", 1

    text = await extract_with_mistral_vision(content, filename)

    if not text:
        return "[Image content - extraction failed]", 1

    return text, 1


async def extract_text_from_txt(content: bytes) -> Tuple[str, int]:
    """
    Extract text from plain text file.

    Returns:
        Tuple of (text, estimated_page_count)
    """
    try:
        # Try UTF-8 first, then fallback
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        word_count = len(text.split())
        page_count = max(1, word_count // 500)

        logger.info(f"Extracted {len(text)} chars from TXT (~{page_count} pages)")
        return text, page_count

    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        return "", 0


async def process_file(filename: str, content: bytes) -> dict:
    """
    Process an uploaded file and extract text.

    Args:
        filename: Name of the file
        content: File content as bytes

    Returns:
        Dict with filename, file_type, extracted_text, word_count, page_count
    """
    filename_lower = filename.lower()

    # Determine file type and extract
    if filename_lower.endswith(".pdf"):
        text, page_count = await extract_text_from_pdf(content, filename)
        file_type = "pdf"
    elif filename_lower.endswith((".docx", ".doc")):
        text, page_count = await extract_text_from_docx(content)
        file_type = "docx"
    elif filename_lower.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp")):
        text, page_count = await extract_text_from_image(content, filename)
        file_type = "image"
    elif filename_lower.endswith(".txt"):
        text, page_count = await extract_text_from_txt(content)
        file_type = "txt"
    else:
        logger.warning(f"Unknown file type: {filename}")
        text = ""
        page_count = 0
        file_type = "unknown"

    word_count = len(text.split()) if text else 0

    return {
        "filename": filename,
        "file_type": file_type,
        "extracted_text": text,
        "word_count": word_count,
        "page_count": page_count,
        "file_size": len(content),
    }


async def process_files(files: list) -> list:
    """
    Process multiple uploaded files.

    Args:
        files: List of (filename, content) tuples or UploadFile objects

    Returns:
        List of processed file dicts
    """
    results = []

    for file_data in files:
        if isinstance(file_data, tuple):
            filename, content = file_data
        elif hasattr(file_data, "filename"):
            # FastAPI UploadFile
            filename = file_data.filename
            content = await file_data.read()
        else:
            continue

        result = await process_file(filename, content)
        results.append(result)

    logger.info(f"Processed {len(results)} files")
    return results
